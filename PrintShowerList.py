import docx
import json
import datetime
td.strftime("%#I:%m %p")
d=docx.Document('ShowerSheet.docx')
a=json.load(open('20230922ShowerList.tsv'))
t=d.tables[0]
for j in range(0,len(a)):
	aa=a[j]
	c=t.rows[j+1].cells
	c[1].paragraphs[0].runs[0].text=f"{aa['showerNum']}"
	c[2].paragraphs[0].runs[0].text="X" if aa['outTime']>0 else ""
	c[3].paragraphs[0].runs[0].text=aa['name']
	c[4].paragraphs[0].runs[0].text=aa['birthday']
	c[5].paragraphs[0].runs[0].text=datetime.datetime.fromtimestamp(aa['inTime']/1000).strftime("%#I:%m %p")
	
d.save('ShowerList.docx')

